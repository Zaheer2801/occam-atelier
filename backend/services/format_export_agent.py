"""
Converts ResumeContent into all output formats:
  - PDF ATS-clean (WeasyPrint, no images)
  - PDF designed (WeasyPrint, with headshot if available)
  - DOCX (python-docx)
  - Plain text (for ATS copy-paste)
"""
from __future__ import annotations

import os
import uuid
from pathlib import Path
from typing import Optional

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML as WP_HTML
from docx import Document
from docx.shared import Pt, Inches

from models.schemas import ProfileJSON, ResumeContent, FormatOutputs

BASE_DIR = Path(__file__).parent.parent
TEMPLATE_DIR = BASE_DIR / "templates"
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "./outputs"))

_jinja = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))


def _ensure_output(resume_id: str) -> Path:
    d = OUTPUT_DIR / resume_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _render_pdf_ats(
    profile: ProfileJSON,
    content: ResumeContent,
    out_dir: Path,
) -> str:
    tmpl = _jinja.get_template("resume_ats.html")
    html = tmpl.render(profile=profile, content=content)
    path = out_dir / "resume_ats.pdf"
    WP_HTML(string=html, base_url=str(BASE_DIR)).write_pdf(str(path))
    (out_dir / "resume_ats.html").write_text(html, encoding="utf-8")
    return str(path)


def _render_pdf_designed(
    profile: ProfileJSON,
    content: ResumeContent,
    out_dir: Path,
    target_role: Optional[str] = None,
    headshot_path: Optional[str] = None,
) -> str:
    import base64 as _b64
    headshot_src = None
    if headshot_path:
        try:
            data = Path(headshot_path).read_bytes()
            headshot_src = f"data:image/jpeg;base64,{_b64.b64encode(data).decode()}"
        except Exception:
            pass

    tmpl = _jinja.get_template("resume_designed.html")
    html = tmpl.render(
        profile=profile,
        content=content,
        target_role=target_role,
        headshot_src=headshot_src,
    )
    path = out_dir / "resume_designed.pdf"
    WP_HTML(string=html, base_url=str(BASE_DIR)).write_pdf(str(path))
    (out_dir / "resume_designed.html").write_text(html, encoding="utf-8")
    return str(path)


def _render_pdf_modern(
    profile: ProfileJSON,
    content: ResumeContent,
    out_dir: Path,
    target_role: Optional[str] = None,
) -> str:
    tmpl = _jinja.get_template("resume_modern.html")
    html = tmpl.render(profile=profile, content=content, target_role=target_role)
    path = out_dir / "resume_modern.pdf"
    WP_HTML(string=html, base_url=str(BASE_DIR)).write_pdf(str(path))
    (out_dir / "resume_modern.html").write_text(html, encoding="utf-8")
    return str(path)


def _render_docx(
    profile: ProfileJSON,
    content: ResumeContent,
    out_dir: Path,
) -> str:
    doc = Document()

    # Name
    name_para = doc.add_paragraph()
    run = name_para.add_run(profile.full_name)
    run.bold = True
    run.font.size = Pt(18)

    # Contact line
    contact_parts = filter(None, [profile.email, profile.phone, profile.location, profile.linkedin])
    doc.add_paragraph(" · ".join(contact_parts)).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    def section(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    if content.summary:
        section("Summary")
        doc.add_paragraph(content.summary).paragraph_format.space_after = Pt(6)

    if content.skills:
        section("Skills")
        doc.add_paragraph(", ".join(content.skills)).paragraph_format.space_after = Pt(6)

    if content.experience:
        section("Experience")
        for exp in content.experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('company', '')}  ").bold = True
            p.add_run(f"{exp.get('start_date', '')} – {exp.get('end_date') or 'Present'}")
            p.paragraph_format.space_after = Pt(0)
            doc.add_paragraph(exp.get("title", "")).runs[0].italic = True
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet").paragraph_format.space_after = Pt(1)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if content.education:
        section("Education")
        for edu in content.education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('institution', '')}").bold = True
            year = edu.get("graduation_year")
            suffix = f"  {year}" if year else ""
            p.add_run(f" — {edu.get('degree', '')}, {edu.get('field', '')}{suffix}")

    path = out_dir / "resume.docx"
    doc.save(str(path))
    return str(path)


def _render_txt(profile: ProfileJSON, content: ResumeContent) -> str:
    lines: list[str] = []
    lines.append(profile.full_name.upper())
    contact = " | ".join(filter(None, [profile.email, profile.phone, profile.location]))
    if contact:
        lines.append(contact)
    lines.append("")

    if content.summary:
        lines += ["SUMMARY", "-" * 40, content.summary, ""]

    if content.skills:
        lines += ["SKILLS", "-" * 40, ", ".join(content.skills), ""]

    if content.experience:
        lines.append("EXPERIENCE")
        lines.append("-" * 40)
        for exp in content.experience:
            lines.append(f"{exp.get('company', '')} | {exp.get('title', '')} | "
                         f"{exp.get('start_date', '')} – {exp.get('end_date') or 'Present'}")
            for bullet in exp.get("bullets", []):
                lines.append(f"  • {bullet}")
            lines.append("")

    if content.education:
        lines.append("EDUCATION")
        lines.append("-" * 40)
        for edu in content.education:
            year = f" ({edu.get('graduation_year')})" if edu.get("graduation_year") else ""
            lines.append(f"{edu.get('institution', '')} — {edu.get('degree', '')}, {edu.get('field', '')}{year}")

    return "\n".join(lines)


async def export_all(
    resume_id: str,
    profile: ProfileJSON,
    content: ResumeContent,
    target_role: Optional[str] = None,
    headshot_path: Optional[str] = None,
) -> FormatOutputs:
    out_dir = _ensure_output(resume_id)

    pdf_ats = _render_pdf_ats(profile, content, out_dir)
    pdf_designed = _render_pdf_designed(profile, content, out_dir, target_role, headshot_path)
    pdf_modern = _render_pdf_modern(profile, content, out_dir, target_role)
    docx_path = _render_docx(profile, content, out_dir)
    txt = _render_txt(profile, content)

    txt_path = out_dir / "resume.txt"
    txt_path.write_text(txt, encoding="utf-8")

    return FormatOutputs(
        pdf_ats_path=pdf_ats,
        pdf_designed_path=pdf_designed,
        pdf_modern_path=pdf_modern,
        docx_path=docx_path,
        txt=txt,
    )
